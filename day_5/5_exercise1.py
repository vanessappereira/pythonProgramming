# exportar os dados armazenados no array para pdf, word e excel
import pickle
from docx import Document
from fpdf import FPDF
from openpyxl import Workbook

lista_pokemon = ["Bulbassaur", "Pikachu", "Machamp", "Vitrebeel"]


def adicionar_pokemon():
    try:
        nome_pokemon = input("Introduza o nome do Pokémon:")
        if nome_pokemon == "":
            raise ValueError
        else:
            lista_pokemon.append(nome_pokemon)
            print(f"{nome_pokemon} adicionado com sucesso.")

    except ValueError:
        print("Resposta inválida.")


def listar_pokemons():
    if not lista_pokemon:
        print("A lista encontra-se vazia! ")
    else:
        print("A Lista contém os seguintes Pokémons: ")
        for i, (nome_pokemon) in enumerate(zip(lista_pokemon), start=1):
            print(f"{i} - {nome_pokemon}")


def remover_pokemon():
    if not lista_pokemon:
        print("A lista encontra-se vazia! ")
    else:
        listar_pokemons()
        try:
            id = int(input("Introduza o id do Pokémon a remover: "))
            if 1 <= id <= len(lista_pokemon):
                pokemon_removido = lista_pokemon.pop(id - 1)
                print(f"{pokemon_removido} eliminado com sucesso! ")
            else:
                print("Id Incorreto! Tente de novo! ")
        except ValueError:
            print("O Id deve ser um nº inteiro! ")


def gerar_ficheiro_pickle():
    global lista_pokemon
    if not lista_pokemon:
        print("A lista encontra-se vazia! ")
    else:
        try:
            with open("lista_pokemon.pickle", "wb") as ficheiro:
                pickle.dump((lista_pokemon), ficheiro)
            print("Dados gerados com sucesso! ")
        except IOError:
            print("Erro ao guardar ficheiro!")


def abrir_ficheiro_pickle():
    global lista_pokemon
    try:
        with open("lista_pokemon.pickle", "rb") as ficheiro:
            lista_pokemon = pickle.load(ficheiro)
        print("Dados Carregados com Sucesso!")
    except FileNotFoundError:
        print(
            "Ficheiro não encontrado! Verifique se o ficheiro consta na diretoria, ou que o nome está bem indicado! "
        )
    except IOError:
        print("Erro ao abrir o ficheiro! ")


def gerar_docx():
    if not lista_pokemon:
        print("A lista encontra-se vazia! ")
    else:
        doc = Document()
        doc.add_heading("Lista de Pokémon", 0)

        for nome_pokemon in zip(lista_pokemon):
            doc.add_paragraph(f"{nome_pokemon} ")
        doc.save("lista_pokemon.docx")
        print("Ficheiro word gerado com sucesso!")


def gerar_pdf():
    if not lista_pokemon:
        print("A lista encontra-se vazia! ")
    else:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=14)

        pdf.cell(200, 10, txt="Lista de Pokémons", ln=True, align="C")

        for nome_pokemon in zip(lista_pokemon):
            pdf.cell(200, 10, txt=f"{nome_pokemon} ", ln=True)

        pdf.output("lista_pokemon.pdf")
        print("Ficheiro PDF gerado com sucesso!")


def gerar_excel():
    if not lista_pokemon:
        print("A lista encontra-se vazia! ")
    else:
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Lista de Pokémons"

        sheet.append(["Título"])

        for nome_pokemon in zip(lista_pokemon):
            sheet.append([nome_pokemon])

        workbook.save("lista_pokemon.xlsx")
        print("Ficheiro Exportado para Excel com sucesso! ")


while True:
    print("Bem-Vindo(a)")
    print("1. Adicionar Pokémons na lista ")
    print("2. Listar Pokémons da lista ")
    print("3. Remover Pokémon da lista ")
    print("5. Gerar Ficheiro Pickle ")
    print("6. Abrir Ficheiro Pickle")
    print("7. Gerar Ficheiro Word")
    print("8. Gerar Ficheiro PDF")
    print("9. Gerar Ficheiro Excel")
    print("10. Sair")

    opcao = input("Introduza a opção a executar: ")

    if opcao == "1":
        adicionar_pokemon()
    elif opcao == "2":
        listar_pokemons()
    elif opcao == "3":
        remover_pokemon()
    elif opcao == "5":
        gerar_ficheiro_pickle()
    elif opcao == "6":
        abrir_ficheiro_pickle()
    elif opcao == "7":
        gerar_docx()
    elif opcao == "8":
        gerar_pdf()
    elif opcao == "9":
        gerar_excel()
    elif opcao == "10":
        print("Programa terminado. Obrigado!")
        break
    else:
        print("Opção inválida!")
